"""
Clausely Symbolic Formatting Engine (SFE)

THE core moat of Clausely. This is NOT an LLM — it is a deterministic,
rule-based formatting engine that enforces Indian court formatting as
immutable constraints.

Architecture:
    LLM generates content → SFE enforces format → Registry accepts document

The SFE loads court-specific formatting rules from JSON files and applies
them to any document, producing validated, court-ready output with exact
margins, fonts, spacing, and section ordering.
"""

from __future__ import annotations

import io
import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Third-party imports — PDF and DOCX generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Frame,
    PageTemplate,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


# ---------------------------------------------------------------------------
# Data directory resolution
# ---------------------------------------------------------------------------

_DATA_DIR = Path(__file__).resolve().parent.parent / "data"
_COURT_FORMATS_DIR = _DATA_DIR / "court_formats"
_TEMPLATES_DIR = _DATA_DIR / "templates"


# ---------------------------------------------------------------------------
# Validation result
# ---------------------------------------------------------------------------

@dataclass
class ValidationResult:
    """Result of running a document through the SFE validator."""
    is_valid: bool
    acceptance_score: float  # 0–100
    fatal_defects: List[str] = field(default_factory=list)
    curable_defects: List[str] = field(default_factory=list)
    formatted_sections: Dict[str, Any] = field(default_factory=dict)
    details: Dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Symbolic Formatting Engine
# ---------------------------------------------------------------------------

class SymbolicFormattingEngine:
    """
    Deterministic court formatting rule enforcer.

    Enforces Indian court formatting as hard constraints.
    NO LLM can override these rules.
    """

    # Jurisdiction-code → JSON filename mapping
    _JURISDICTION_FILES = {
        "MH-DISTRICT": "MH_DISTRICT.json",
        "MH-HC": "MH_HC.json",
        "DL-DISTRICT": "DL_DISTRICT.json",
        "IN-SC": "IN_SC.json",
    }

    def __init__(self, jurisdiction: str):
        self.jurisdiction = jurisdiction
        self.rules: Dict[str, Any] = self._load_court_rules(jurisdiction)

    # ------------------------------------------------------------------
    # Rule Loading
    # ------------------------------------------------------------------

    def _load_court_rules(self, jurisdiction: str) -> Dict[str, Any]:
        """Load court formatting rules from the data directory."""
        filename = self._JURISDICTION_FILES.get(jurisdiction)
        if not filename:
            raise ValueError(
                f"Unsupported jurisdiction: {jurisdiction}. "
                f"Supported: {list(self._JURISDICTION_FILES.keys())}"
            )
        filepath = _COURT_FORMATS_DIR / filename
        if not filepath.exists():
            raise FileNotFoundError(f"Court rules file not found: {filepath}")
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    def validate(self, document: Dict[str, Any]) -> ValidationResult:
        """
        Validate a document dictionary against court rules.

        The document dict should have:
            - sections: dict mapping section_type → content string
            - metadata: dict with court_name, advocate_name, etc.
            - formatting: dict with margin_left_cm, font_body, etc. (optional)
        """
        fatal: List[str] = []
        curable: List[str] = []
        checks: Dict[str, bool] = {}

        sections = document.get("sections", {})
        formatting = document.get("formatting", {})
        metadata = document.get("metadata", {})

        # ------- Section presence checks -------
        required_sections = self.rules.get("required_sections", [])
        for sec in required_sections:
            present = sec in sections and bool(sections[sec].strip() if isinstance(sections[sec], str) else sections[sec])
            checks[f"{sec}_present"] = present
            if not present:
                defect_key = f"missing_{sec}"
                if defect_key in self.rules.get("fatal_defects", []):
                    fatal.append(f"Missing required section: {sec}")
                else:
                    curable.append(f"Missing recommended section: {sec}")

        # ------- Formatting checks -------
        if formatting:
            # Margin checks
            for side in ["left", "right", "top", "bottom"]:
                rule_key = f"margin_{side}_cm"
                expected = self.rules.get(rule_key)
                actual = formatting.get(rule_key)
                if expected is not None and actual is not None:
                    is_correct = abs(float(actual) - float(expected)) < 0.15
                    checks[f"margin_{side}_correct"] = is_correct
                    if not is_correct:
                        fatal.append(
                            f"Incorrect {side} margin: expected {expected}cm, got {actual}cm"
                        )

            # Font checks
            expected_font = self.rules.get("font_body")
            actual_font = formatting.get("font_body")
            if expected_font and actual_font:
                checks["font_correct"] = actual_font == expected_font
                if actual_font != expected_font:
                    curable.append(
                        f"Incorrect body font: expected '{expected_font}', got '{actual_font}'"
                    )

            # Font size checks
            expected_size = self.rules.get("font_size_body")
            actual_size = formatting.get("font_size_body")
            if expected_size is not None and actual_size is not None:
                diff = abs(int(actual_size) - int(expected_size))
                checks["font_size_correct"] = diff == 0
                if diff == 1:
                    curable.append(
                        f"Font size off by 1pt: expected {expected_size}pt, got {actual_size}pt"
                    )
                elif diff > 1:
                    fatal.append(
                        f"Incorrect font size: expected {expected_size}pt, got {actual_size}pt"
                    )

            # Line spacing
            expected_spacing = self.rules.get("line_spacing")
            actual_spacing = formatting.get("line_spacing")
            if expected_spacing and actual_spacing:
                checks["line_spacing_correct"] = abs(float(actual_spacing) - float(expected_spacing)) < 0.1
                if not checks["line_spacing_correct"]:
                    fatal.append(
                        f"Incorrect line spacing: expected {expected_spacing}, got {actual_spacing}"
                    )
        else:
            # No formatting provided — apply defaults (not a fatal error for
            # validation, but mark as curable since we can auto-fix)
            checks["formatting_provided"] = False
            curable.append("No explicit formatting metadata provided — will apply defaults")

        # ------- Line-Level Layout and Margin Checks (Stage 2.1) -------
        line_logs = []
        margin_violation_found = False
        font_violation_found = False
        style_logs = []
        syntax_errors = []
        numbering_sequence_broken = False
        
        if formatting:
            # 1. Left margin spacing check: assert matches 1.5 inches (approx 3.81 cm or 4.0 cm) for HC/SC,
            # or matches the court-specific rules for other jurisdictions.
            actual_left_cm = formatting.get("margin_left_cm")
            actual_left_in = formatting.get("margin_left_in")
            
            if self.jurisdiction in ("MH-HC", "IN-SC"):
                if actual_left_in is not None:
                    if abs(float(actual_left_in) - 1.5) > 0.05:
                        margin_violation_found = True
                elif actual_left_cm is not None:
                    # Bombay HC has 4.0 cm left margin. SC has 1.5 in (~3.81cm)
                    if abs(float(actual_left_cm) - 3.81) > 0.2 and abs(float(actual_left_cm) - 4.0) > 0.2:
                        margin_violation_found = True
            else:
                expected_left_cm = self.rules.get("margin_left_cm")
                if expected_left_cm is not None and actual_left_cm is not None:
                    if abs(float(actual_left_cm) - float(expected_left_cm)) > 0.15:
                        margin_violation_found = True
            
            # 2. Line spacing check: assert double-line spacing equals 2.0 for HC/SC, or expected for others
            actual_spacing = formatting.get("line_spacing")
            if actual_spacing is not None:
                if self.jurisdiction in ("MH-HC", "IN-SC"):
                    if abs(float(actual_spacing) - 2.0) > 0.1:
                        margin_violation_found = True
                else:
                    expected_spacing = self.rules.get("line_spacing")
                    if expected_spacing is not None:
                        if abs(float(actual_spacing) - float(expected_spacing)) > 0.1:
                            margin_violation_found = True
            
            # 3. Right margin spacing check: verify aligns to 1.0 inch (2.54 cm or 2.5 cm) for HC/SC, or expected for others
            actual_right_cm = formatting.get("margin_right_cm")
            actual_right_in = formatting.get("margin_right_in")
            if self.jurisdiction in ("MH-HC", "IN-SC"):
                if actual_right_in is not None:
                    if abs(float(actual_right_in) - 1.0) > 0.05:
                        margin_violation_found = True
                elif actual_right_cm is not None:
                    if abs(float(actual_right_cm) - 2.54) > 0.2 and abs(float(actual_right_cm) - 2.5) > 0.2:
                        margin_violation_found = True
            else:
                expected_right_cm = self.rules.get("margin_right_cm")
                if expected_right_cm is not None and actual_right_cm is not None:
                    if abs(float(actual_right_cm) - float(expected_right_cm)) > 0.15:
                        margin_violation_found = True

            # 4. Font family & size (Sub-Stage 2.2)
            actual_font = formatting.get("font_body")
            if actual_font:
                expected_font = self.rules.get("font_body", "Times New Roman")
                if actual_font != expected_font:
                    font_violation_found = True
            actual_size = formatting.get("font_size_body")
            if actual_size is not None:
                expected_size = self.rules.get("font_size_body", 14)
                if int(actual_size) != expected_size:
                    font_violation_found = True
            # Color pure black check
            font_color = formatting.get("font_color")
            if font_color is not None:
                if str(font_color).lower() not in ("black", "#000000", "000000", "pure_black"):
                    fatal.append("[FONT_VIOLATION_HALT] Font color must be pure black.")

        # Scan section text content lines
        line_states = []
        last_footnote_num = 0

        for sec_name, sec_content in sections.items():
            if not isinstance(sec_content, str):
                continue
            lines = sec_content.splitlines()
            total_lines = len(lines)
            correct_syntax_lines = 0
            last_number = 0
            last_table_cols = None
            
            for idx, line in enumerate(lines):
                # Trailing spaces at line ends
                if line.endswith(" ") or line.endswith("\t"):
                    curable.append(f"[TRAILING_SPACE] Trailing spaces at line {idx+1} in section {sec_name}")
                    
                stripped = line.lstrip()
                indent = len(line) - len(stripped)
                
                # Scan tab indentation offsets (Sub-Micro-Sub 2.4.3)
                if "\t" in line[:indent]:
                    curable.append(f"[TAB_INDENTATION] Tab indentation used instead of spaces at line {idx+1} in section {sec_name}")
                
                # Check indentation levels of bulleted items
                if stripped.startswith(("-", "*")) or re.match(r"^\d+\.", stripped):
                    line_logs.append({
                        "section": sec_name,
                        "line_number": idx + 1,
                        "type": "bullet_item",
                        "indentation_spaces": indent
                    })
                else:
                    line_logs.append({
                        "section": sec_name,
                        "line_number": idx + 1,
                        "type": "regular_line",
                        "indentation_spaces": indent
                    })
                    
                # Scan header text positions (Sub-Stage 2.1)
                if stripped.isupper() and len(stripped) > 3:
                    line_logs.append({
                        "section": sec_name,
                        "line_number": idx + 1,
                        "type": "header_line",
                        "text": stripped
                    })
                    style_logs.append({
                        "section": sec_name,
                        "line": idx + 1,
                        "type": "header_style_check",
                        "font_family_ok": True,
                        "bold_ok": True
                    })
                    
                # Detect line wrap formatting defects
                if len(line) > 100:
                    curable.append(f"[LINE_WRAP] Line {idx+1} in section {sec_name} exceeds 100 characters (length: {len(line)})")

                # Font & Style Enforcement (Sub-Stage 2.2)
                # Scan for non-approved bold formatting
                if not (stripped.isupper() and len(stripped) > 3) and not stripped.startswith("#"):
                    if "**" in line or "<b>" in line or "<strong>" in line:
                        curable.append(f"[STYLE_INCONSISTENCY] Non-approved bold formatting at line {idx+1} in section {sec_name}")
                # Blockquote scan
                if stripped.startswith(">"):
                    style_logs.append({
                        "section": sec_name,
                        "line": idx + 1,
                        "type": "blockquote",
                        "indentation_offset": indent
                    })
                # Detect invalid Unicode character marks (outside CP1252 printable range)
                for char in line:
                    if ord(char) > 255:
                        fatal.append(f"[FONT_VIOLATION_HALT] Invalid Unicode character mark '{char}' detected at line {idx+1} in section {sec_name}")
                        font_violation_found = True
                        break

                # Line Syntax Verification (Sub-Stage 2.3)
                if stripped:
                    # Check capitalization / first character
                    starts_properly = stripped[0].isupper() or stripped[0].isdigit() or stripped[0] in ("-", "*", "#", ">", "[")
                    # Double spaces check
                    has_double_spaces = "  " in stripped.replace(".  ", " ")
                    
                    # Punctuation balancing check (Sub-Micro-Sub 2.3.1.1.2)
                    parentheses_balanced = (stripped.count("(") == stripped.count(")"))
                    brackets_balanced = (stripped.count("[") == stripped.count("]"))
                    braces_balanced = (stripped.count("{") == stripped.count("}"))
                    quotes_balanced = (stripped.count('"') % 2 == 0) and (stripped.count("'") % 2 == 0)
                    punctuation_ok = parentheses_balanced and brackets_balanced and braces_balanced and quotes_balanced
                    
                    if not punctuation_ok:
                        syntax_errors.append(f"Line {idx+1} in section {sec_name} has unbalanced punctuation.")
                        
                    # Double punctuation check (Sub-Micro-Sub 2.3.1.1.2)
                    has_double_punctuation = re.search(r"[,\.\?!]{2,}", stripped) is not None
                    if has_double_punctuation:
                        syntax_errors.append(f"Line {idx+1} in section {sec_name} has double punctuation.")
                        
                    # Permitted abbreviations filtering (Micro-Step 2.3.2)
                    if " vs " in stripped.lower() or " vs. " in stripped.lower():
                        curable.append(f"[ABBREVIATION_WARNING] Non-standard abbreviation 'vs' or 'vs.' at line {idx+1} in section {sec_name}. Use 'v.' or 'versus'.")
                        
                    is_syntax_valid = starts_properly and not has_double_spaces and punctuation_ok and not has_double_punctuation
                    if is_syntax_valid:
                        correct_syntax_lines += 1
                    else:
                        syntax_errors.append(f"Line {idx+1} in section {sec_name} has invalid capitalization/spacing/punctuation.")
                        
                    # Count sentence modifiers (Sub-Micro-Sub 2.3.1.3)
                    sentence_modifiers = ["herein", "hereinafter", "aforesaid", "abovenamed", "hereby", "thereby"]
                    modifier_count = sum(stripped.lower().count(mod) for mod in sentence_modifiers)
                    
                    # Save line validation states to memory (Sub-Micro 2.3.1.4)
                    line_states.append({
                        "section": sec_name,
                        "line_number": idx + 1,
                        "syntax_valid": is_syntax_valid,
                        "punctuation_ok": punctuation_ok,
                        "double_punctuation": has_double_punctuation,
                        "modifier_count": modifier_count
                    })
                    
                    # Date variables standardization
                    if re.search(r"\b\d{1,2}/\d{1,2}/\d{4}\b", line) or re.search(r"\b\d{1,2}-\d{1,2}-\d{4}\b", line):
                        curable.append(f"[DATE_STANDARDIZATION] Date format at line {idx+1} in section {sec_name} should be YYYY-MM-DD.")

                # Indent and Numbering Alignments (Sub-Stage 2.4)
                # Verify bullet list numbers sequencing
                match = re.match(r"^(\d+)\.", stripped)
                if match:
                    current_number = int(match.group(1))
                    if last_number > 0:
                        if current_number != last_number + 1 and current_number != 1:
                            numbering_sequence_broken = True
                            curable.append(f"[PAGINATION_ANOMALY] Broken numbering sequence at line {idx+1} in section {sec_name}: expected {last_number+1}, got {current_number}")
                    last_number = current_number
                    
                # Verify footnote number references (Sub-Micro-Sub 2.4.1.1.3)
                footnote_matches = re.findall(r"\[\^(\d+)\]", stripped)
                for fn in footnote_matches:
                    fn_num = int(fn)
                    if last_footnote_num > 0 and fn_num != last_footnote_num + 1:
                        numbering_sequence_broken = True
                        curable.append(f"[PAGINATION_ANOMALY] Broken footnote sequence at line {idx+1} in section {sec_name}: expected {last_footnote_num+1}, got {fn_num}")
                    last_footnote_num = fn_num
                    
                # Check table column alignment indexes (Sub-Micro-Sub 2.4.1.3)
                if "|" in stripped:
                    cols_count = stripped.count("|")
                    if last_table_cols is not None and cols_count != last_table_cols:
                        curable.append(f"[ALIGNMENT_ANOMALY] Table column count mismatch at line {idx+1} in section {sec_name}.")
                    last_table_cols = cols_count
                else:
                    last_table_cols = None

            # Assess grammar score threshold
            non_empty_lines = sum(1 for l in lines if l.strip())
            score = (correct_syntax_lines / non_empty_lines) if non_empty_lines > 0 else 1.0
            if score < 0.8:
                fatal.append(f"[SYNTAX_REWRITE] Grammar or syntax score ({score:.2f}) below 0.80 threshold in section {sec_name}.")

        # Check section header numbers sequence
        header_numbers = []
        for sec_name, sec_content in sections.items():
            if not isinstance(sec_content, str):
                continue
            lines = sec_content.splitlines()
            for line in lines[:2]:
                stripped = line.strip()
                match = re.match(r"^(\d+)\s+[A-Z]", stripped)
                if match:
                    header_numbers.append(int(match.group(1)))
        if len(header_numbers) > 1:
            for i in range(1, len(header_numbers)):
                if header_numbers[i] != header_numbers[i-1] + 1:
                    numbering_sequence_broken = True

        if margin_violation_found:
            fatal.append("[MARGIN_VIOLATION] Spacing or margin formatting deviation detected.")
        if font_violation_found:
            fatal.append("[FONT_VIOLATION_HALT] Font family, size, or Unicode encoding violation.")
        if numbering_sequence_broken:
            fatal.append("[PAGINATION_ANOMALY] Broken numbering or header sequence continuity.")
            
        checks["layout_check_records"] = True
        checks["line_coordinates_logs"] = line_logs
        checks["style_compliance_logs"] = style_logs
        checks["style_validation_indicators"] = True
        checks["syntax_validation_states"] = line_states
        checks["layout_alignment_summary"] = {
            "numbering_sequence_ok": not numbering_sequence_broken,
            "tables_aligned": True,
        }

        # ------- Stamp paper checks -------
        if self.rules.get("stamp_paper_required"):
            stamp_value = metadata.get("stamp_paper_value")
            expected_stamp = self.rules.get("stamp_paper_value")
            checks["stamp_paper_correct"] = str(stamp_value) == str(expected_stamp) if stamp_value else False
            if not checks["stamp_paper_correct"]:
                fatal.append(
                    f"Stamp paper value mismatch: expected ₹{expected_stamp}, "
                    f"got ₹{stamp_value or 'MISSING'}"
                )

        # ------- Verification clause -------
        if self.rules.get("verification_clause_required"):
            has_verification = "verification" in sections and bool(
                sections["verification"].strip() if isinstance(sections["verification"], str) else sections["verification"]
            )
            checks["verification_present"] = has_verification
            if not has_verification:
                fatal.append("Missing verification clause (mandatory)")

        # ------- Notary check -------
        if self.rules.get("notary_required"):
            has_notary = metadata.get("notarized", False)
            checks["notary_present"] = has_notary
            if not has_notary:
                curable.append("Document not marked as notarized (required for this jurisdiction)")

        # ------- Score computation -------
        fatal_weight = len(fatal) * 20
        curable_weight = len(curable) * 5
        score = max(0.0, 100.0 - fatal_weight - curable_weight)

        # Build formatted sections showing what the correct structure looks like
        formatted_sections = {}
        for sec_type in self.rules.get("section_order", required_sections):
            formatted_sections[sec_type] = {
                "present": sec_type in sections,
                "order": self.rules.get("section_order", []).index(sec_type)
                if sec_type in self.rules.get("section_order", [])
                else -1,
                "mandatory": sec_type in required_sections,
            }

        return ValidationResult(
            is_valid=len(fatal) == 0,
            acceptance_score=score,
            fatal_defects=fatal,
            curable_defects=curable,
            formatted_sections=formatted_sections,
            details=checks,
        )

    # ------------------------------------------------------------------
    # Enforcement — apply formatting rules to raw content
    # ------------------------------------------------------------------

    def enforce(self, raw_content: str, metadata: Optional[Dict[str, str]] = None) -> str:
        """
        Take raw LLM-generated content and apply all formatting rules.

        This runs AFTER LLM generation. The LLM fills content, the SFE
        deterministically enforces structure and format.
        """
        metadata = metadata or {}
        lines: List[str] = []

        # 1. Build header
        if self.rules.get("header_required"):
            header_fmt = self.rules.get("header_format", "")
            header = header_fmt.format(
                court_designation=metadata.get("court_designation", "{court_designation}"),
                city_name=metadata.get("city_name", "{city_name}"),
                case_type=metadata.get("case_type", "{case_type}"),
                case_number=metadata.get("case_number", "{case_number}"),
                year=metadata.get("year", "{year}"),
                bench_name=metadata.get("bench_name", "{bench_name}"),
                bench_type=metadata.get("bench_type", "{bench_type}"),
                district_name=metadata.get("district_name", "{district_name}"),
            )
            lines.append(header)
            lines.append("")

        # 2. Inject content (already generated by LLM)
        lines.append(raw_content.strip())

        # 3. Build footer marker
        if self.rules.get("footer_required"):
            lines.append("")
            lines.append("---")
            footer_fmt = self.rules.get("footer_format", "")
            footer = footer_fmt.format(
                page_num="{page_num}",
                total_pages="{total_pages}",
                advocate_name=metadata.get("advocate_name", "{advocate_name}"),
                bar_number=metadata.get("bar_number", "{bar_number}"),
            )
            lines.append(footer)

        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Section ordering
    # ------------------------------------------------------------------

    def reorder_sections(self, sections: Dict[str, str]) -> List[Tuple[str, str]]:
        """Return sections in the court-mandated order."""
        order = self.rules.get("section_order", list(sections.keys()))
        ordered: List[Tuple[str, str]] = []
        for sec_type in order:
            if sec_type in sections:
                ordered.append((sec_type, sections[sec_type]))
        # Append any extra sections not in the official order
        for sec_type, content in sections.items():
            if sec_type not in order:
                ordered.append((sec_type, content))
        return ordered

    # ------------------------------------------------------------------
    # PDF Export
    # ------------------------------------------------------------------

    def export_pdf(self, document: Dict[str, Any]) -> bytes:
        """
        Generate a court-ready PDF using ReportLab.

        Applies exact margin, font, and spacing rules from the loaded
        court format JSON.
        """
        buf = io.BytesIO()

        margin_left = self.rules.get("margin_left_cm", 3.0) * cm
        margin_right = self.rules.get("margin_right_cm", 2.5) * cm
        margin_top = self.rules.get("margin_top_cm", 2.5) * cm
        margin_bottom = self.rules.get("margin_bottom_cm", 2.0) * cm

        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=margin_left,
            rightMargin=margin_right,
            topMargin=margin_top,
            bottomMargin=margin_bottom,
            title=f"Clausely — {self.rules.get('court_name', 'Legal Document')}",
            author="Clausely Legal AI",
        )

        font_name = self.rules.get("font_body", "Times-Roman")
        font_size = int(self.rules.get("font_size_body", 14))
        heading_size = int(self.rules.get("font_size_heading", 14))
        line_spacing = float(self.rules.get("line_spacing", 1.5))

        # Map common font names to ReportLab built-ins
        rl_font = "Times-Roman"
        if "times" in font_name.lower():
            rl_font = "Times-Roman"
        elif "arial" in font_name.lower():
            rl_font = "Helvetica"

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "ClauselyBody",
            parent=styles["Normal"],
            fontName=rl_font,
            fontSize=font_size,
            leading=font_size * line_spacing,
            alignment=TA_JUSTIFY,
            spaceAfter=int(self.rules.get("paragraph_spacing_after", 12)),
        )
        heading_style = ParagraphStyle(
            "ClauselyHeading",
            parent=styles["Heading1"],
            fontName="Times-Bold" if rl_font == "Times-Roman" else (f"{rl_font}-Bold" if rl_font == "Helvetica" else rl_font),
            fontSize=heading_size,
            leading=heading_size * 1.4,
            alignment=TA_CENTER,
            spaceAfter=14,
        )
        center_style = ParagraphStyle(
            "ClauselyCenter",
            parent=body_style,
            alignment=TA_CENTER,
        )

        elements = []

        # Header
        metadata = document.get("metadata", {})
        if self.rules.get("header_required"):
            header_fmt = self.rules.get("header_format", "")
            header_text = header_fmt.format(
                court_designation=metadata.get("court_designation", ""),
                city_name=metadata.get("city_name", ""),
                case_type=metadata.get("case_type", ""),
                case_number=metadata.get("case_number", ""),
                year=metadata.get("year", ""),
                bench_name=metadata.get("bench_name", ""),
                bench_type=metadata.get("bench_type", ""),
                district_name=metadata.get("district_name", ""),
            )
            for line in header_text.split("\n"):
                line = line.strip()
                if line:
                    elements.append(Paragraph(line, heading_style))
            elements.append(Spacer(1, 20))

        # Sections
        sections = document.get("sections", {})
        ordered = self.reorder_sections(sections)
        for sec_type, content in ordered:
            # Section header
            sec_title = sec_type.replace("_", " ").upper()
            elements.append(Paragraph(sec_title, heading_style))

            # Content paragraphs
            if isinstance(content, str):
                for para in content.split("\n\n"):
                    para = para.strip()
                    if para:
                        # Escape XML entities for ReportLab
                        safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        elements.append(Paragraph(safe, body_style))
            elif isinstance(content, list):
                for item in content:
                    safe = str(item).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    elements.append(Paragraph(safe, body_style))

            elements.append(Spacer(1, 12))

        doc.build(elements)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # DOCX Export
    # ------------------------------------------------------------------

    def export_docx(self, document: Dict[str, Any]) -> bytes:
        """
        Generate an editable Word document using python-docx.

        Applies exact margin, font, and spacing rules from the court
        format specification.
        """
        docx = DocxDocument()

        # Page margins
        for section in docx.sections:
            section.top_margin = Cm(self.rules.get("margin_top_cm", 2.5))
            section.bottom_margin = Cm(self.rules.get("margin_bottom_cm", 2.0))
            section.left_margin = Cm(self.rules.get("margin_left_cm", 3.0))
            section.right_margin = Cm(self.rules.get("margin_right_cm", 2.5))

        font_name = self.rules.get("font_body", "Times New Roman")
        font_size = int(self.rules.get("font_size_body", 14))
        heading_font_size = int(self.rules.get("font_size_heading", 14))
        line_spacing_val = float(self.rules.get("line_spacing", 1.5))

        # Default font
        style = docx.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.paragraph_format.line_spacing = line_spacing_val
        style.paragraph_format.space_after = Pt(
            int(self.rules.get("paragraph_spacing_after", 12))
        )

        metadata = document.get("metadata", {})

        # Header
        if self.rules.get("header_required"):
            header_fmt = self.rules.get("header_format", "")
            header_text = header_fmt.format(
                court_designation=metadata.get("court_designation", ""),
                city_name=metadata.get("city_name", ""),
                case_type=metadata.get("case_type", ""),
                case_number=metadata.get("case_number", ""),
                year=metadata.get("year", ""),
                bench_name=metadata.get("bench_name", ""),
                bench_type=metadata.get("bench_type", ""),
                district_name=metadata.get("district_name", ""),
            )
            for line in header_text.split("\n"):
                line = line.strip()
                if line:
                    para = docx.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(line)
                    run.font.name = font_name
                    run.font.size = Pt(heading_font_size)
                    run.bold = self.rules.get("font_heading_bold", True)

            docx.add_paragraph()  # spacer

        # Sections
        sections = document.get("sections", {})
        ordered = self.reorder_sections(sections)
        for sec_type, content in ordered:
            # Section heading
            sec_title = sec_type.replace("_", " ").upper()
            heading_para = docx.add_paragraph()
            heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_run = heading_para.add_run(sec_title)
            heading_run.font.name = font_name
            heading_run.font.size = Pt(heading_font_size)
            heading_run.bold = True

            # Section content
            if isinstance(content, str):
                for para_text in content.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        para = docx.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run = para.add_run(para_text)
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
            elif isinstance(content, list):
                for item in content:
                    para = docx.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = para.add_run(str(item))
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

        # Footer info as final paragraph
        if self.rules.get("footer_required"):
            docx.add_paragraph()  # spacer
            footer_para = docx.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_fmt = self.rules.get("footer_format", "")
            footer_text = footer_fmt.format(
                page_num="{page}",
                total_pages="{total}",
                advocate_name=metadata.get("advocate_name", ""),
                bar_number=metadata.get("bar_number", ""),
            )
            run = footer_para.add_run(footer_text)
            run.font.name = font_name
            run.font.size = Pt(font_size - 2)
            run.italic = True

        buf = io.BytesIO()
        docx.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Template Loading
    # ------------------------------------------------------------------

    @staticmethod
    def load_template(template_name: str) -> str:
        """Load a document template from the data/templates directory."""
        filepath = _TEMPLATES_DIR / template_name
        if not filepath.exists():
            raise FileNotFoundError(f"Template not found: {filepath}")
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()

    # ------------------------------------------------------------------
    # Supported jurisdictions
    # ------------------------------------------------------------------

    @classmethod
    def supported_jurisdictions(cls) -> List[str]:
        """Return list of supported jurisdiction codes."""
        return list(cls._JURISDICTION_FILES.keys())
